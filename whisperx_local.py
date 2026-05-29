from __future__ import annotations

import argparse
import os
import time
from pathlib import Path

import av
import ctranslate2
import numpy as np
from faster_whisper import BatchedInferencePipeline, WhisperModel
from docx import Document


def choose_compute_type(requested: str) -> str:
    try:
        supported = ctranslate2.get_supported_compute_types("cuda")
    except Exception as exc:
        raise RuntimeError(
            "CUDA is not available for CTranslate2. This script is GPU-only."
        ) from exc

    if requested in supported:
        return requested

    preferred = [
        "float16",
        "int8_float16",
        "bfloat16",
        "int8_bfloat16",
        "float32",
        "int8_float32",
        "int8",
    ]
    for candidate in preferred:
        if candidate in supported:
            return candidate

    return sorted(supported)[0] if supported else "float32"


def write_txt(segments: list[dict], out_file: Path) -> None:
    lines: list[str] = []
    for seg in segments:
        text = seg.get("text", "").strip()
        if not text:
            continue

        speaker = seg.get("speaker")
        if speaker:
            lines.append(f"[{speaker}] {text}")
        else:
            lines.append(text)

    text = "\n".join(lines)
    out_file.write_text(text, encoding="utf-8")


def write_docx(segments: list[dict], out_file: Path, title: str | None = None) -> None:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for seg in segments:
        text = seg.get("text", "").strip()
        if text:
            speaker = seg.get("speaker")
            if speaker:
                p = doc.add_paragraph()
                p.add_run(f"[{speaker}] ").bold = True
                p.add_run(text)
            else:
                doc.add_paragraph(text)

    doc.save(str(out_file))


def get_audio_duration_seconds(audio_file: Path) -> float:
    with av.open(str(audio_file)) as container:
        if container.duration is not None:
            return float(container.duration / av.time_base)

        audio_stream = next((s for s in container.streams if s.type == "audio"), None)
        if (
            audio_stream is not None
            and audio_stream.duration is not None
            and audio_stream.time_base is not None
        ):
            return float(audio_stream.duration * audio_stream.time_base)

    raise RuntimeError(f"Could not read duration from audio file: {audio_file}")


def load_audio_waveform_16k(audio_file: Path) -> np.ndarray:
    chunks: list[np.ndarray] = []
    with av.open(str(audio_file)) as container:
        audio_stream = next((s for s in container.streams if s.type == "audio"), None)
        if audio_stream is None:
            raise RuntimeError(f"No audio stream found in file: {audio_file}")

        resampler = av.audio.resampler.AudioResampler(
            format="fltp",
            layout="mono",
            rate=16000,
        )

        for frame in container.decode(audio_stream):
            resampled = resampler.resample(frame)
            if resampled is None:
                continue
            frames = resampled if isinstance(resampled, list) else [resampled]
            for fr in frames:
                arr = fr.to_ndarray()
                if arr.ndim == 2:
                    arr = arr[0]
                chunks.append(arr.astype(np.float32, copy=False))

    if not chunks:
        return np.zeros(0, dtype=np.float32)
    return np.concatenate(chunks)


def build_clip_timestamps(duration_s: float, chunk_length_s: int) -> list[dict[str, float]]:
    clips: list[dict[str, float]] = []
    start = 0.0
    while start < duration_s:
        end = min(start + float(chunk_length_s), duration_s)
        if end - start <= 0.0:
            break
        clips.append({"start": start, "end": end})
        start = end
    return clips


def batch_candidates(target_batch_size: int) -> list[int]:
    ordered = [64, 48, 40, 32, 28, 24, 20, 16, 12, 10, 8, 6, 4, 2, 1]
    target_batch_size = max(1, int(target_batch_size))
    if target_batch_size not in ordered:
        ordered.insert(0, target_batch_size)
    return [b for b in ordered if b <= target_batch_size]


def apply_preset(args: argparse.Namespace) -> None:
    preset = (args.preset or "").lower()
    if not preset:
        return

    if preset == "fast":
        args.batch_size = 64
        args.beam_size = 1
        args.best_of = 1
        args.condition_on_previous_text = False
        return

    if preset == "balanced":
        args.batch_size = 32
        args.beam_size = 3
        args.best_of = 3
        args.condition_on_previous_text = False
        return

    if preset == "quality":
        args.batch_size = 24
        args.beam_size = 5
        args.best_of = 5
        args.condition_on_previous_text = False
        return

    raise ValueError(f"Invalid preset: {args.preset}")


def build_whisperx_asr_options(
    beam_size: int,
    best_of: int,
    patience: float,
    temperature: float,
    condition_on_previous_text: bool,
) -> dict:
    return {
        "beam_size": max(1, int(beam_size)),
        "best_of": max(1, int(best_of)),
        "patience": float(patience),
        "temperatures": [float(temperature)],
        "condition_on_previous_text": bool(condition_on_previous_text),
        "without_timestamps": True,
        "word_timestamps": False,
    }


def run(
    audio_file: Path,
    output_dir: Path,
    model_name: str,
    language: str | None,
    batch_size: int,
    requested_compute_type: str,
    beam_size: int,
    best_of: int,
    patience: float,
    temperature: float,
    condition_on_previous_text: bool,
    chunk_length: int | None,
    log_progress: bool,
    auto_batch: bool,
    output_format: str,
    docx_title: str | None,
    diarize: bool,
    hf_token: str | None,
    diarization_device: str,
    num_speakers: int | None,
    min_speakers: int | None,
    max_speakers: int | None,
) -> None:
    if not audio_file.exists():
        raise FileNotFoundError(f"Audio file not found: {audio_file}")

    output_dir.mkdir(parents=True, exist_ok=True)

    compute_type = choose_compute_type(requested_compute_type)
    print(f"Config -> ASR device: cuda, compute_type: {compute_type}")
    print(
        "Decode -> "
        f"beam_size={beam_size}, best_of={best_of}, patience={patience}, "
        f"temperature={temperature}, "
        f"condition_on_previous_text={condition_on_previous_text}, "
        f"chunk_length={chunk_length}"
    )
    print(f"Batch -> target batch_size={batch_size}, auto_batch={auto_batch}, diarize={diarize}")

    duration_s = get_audio_duration_seconds(audio_file)
    chosen_chunk = chunk_length or 30
    print(f"Audio -> duration={duration_s:.1f}s")

    last_exc: Exception | None = None
    segments: list[dict] = []
    started = time.perf_counter()
    used_batch_size = None

    if diarize:
        # WhisperX route: ASR on GPU + diarization on CPU/CUDA.
        import torch
        import whisperx
        from whisperx.diarize import DiarizationPipeline, assign_word_speakers

        token = hf_token or os.getenv("HF_TOKEN") or os.getenv("HUGGINGFACE_TOKEN")
        if not token:
            raise RuntimeError(
                "Diarization requires a Hugging Face token. "
                "Pass --hf-token or set HF_TOKEN."
            )

        asr_options = build_whisperx_asr_options(
            beam_size=beam_size,
            best_of=best_of,
            patience=patience,
            temperature=temperature,
            condition_on_previous_text=condition_on_previous_text,
        )
        model = whisperx.load_model(
            model_name,
            "cuda",
            compute_type=compute_type,
            asr_options=asr_options,
            language=language,
            vad_method="silero",
            vad_options={"chunk_size": chosen_chunk},
        )
        audio = load_audio_waveform_16k(audio_file)

        candidates = batch_candidates(batch_size) if auto_batch else [max(1, batch_size)]
        result = None
        for candidate_bs in candidates:
            try:
                print(f"Trying ASR batch_size={candidate_bs}")
                result = model.transcribe(
                    audio,
                    batch_size=candidate_bs,
                    chunk_size=chosen_chunk,
                    print_progress=log_progress,
                )
                used_batch_size = candidate_bs
                print(f"ASR batch_size={candidate_bs} OK")
                break
            except Exception as exc:
                last_exc = exc
                message = str(exc).lower()
                if "out of memory" in message or ("cuda" in message and "memory" in message):
                    print(f"OOM with batch_size={candidate_bs}, trying lower...")
                    continue
                raise

        if result is None:
            if last_exc is not None:
                raise RuntimeError("All tested ASR batch sizes failed.") from last_exc
            raise RuntimeError("Could not run WhisperX transcription.")

        # Free ASR model memory before diarization stage.
        del model
        if torch.cuda.is_available():
            torch.cuda.empty_cache()

        try:
            diarize_model = DiarizationPipeline(
                token=token,
                device=diarization_device,
            )
        except Exception as exc:
            message = str(exc).lower()
            if "gated" in message or "401" in message or "access" in message:
                raise RuntimeError(
                    "Could not access pyannote diarization model. "
                    "Accept terms at https://huggingface.co/pyannote/speaker-diarization-community-1 "
                    "and use a valid --hf-token."
                ) from exc
            raise
        diarize_kwargs: dict[str, int] = {}
        if num_speakers is not None:
            diarize_kwargs["num_speakers"] = num_speakers
        if min_speakers is not None:
            diarize_kwargs["min_speakers"] = min_speakers
        if max_speakers is not None:
            diarize_kwargs["max_speakers"] = max_speakers

        print(
            "Running diarization -> "
            f"device={diarization_device}, constraints={diarize_kwargs or 'auto'}"
        )
        if audio.size == 0:
            raise RuntimeError("Decoded audio is empty; cannot run diarization.")
        try:
            diarize_segments = diarize_model(audio, **diarize_kwargs)
        except Exception as exc:
            message = str(exc).lower()
            if "gated" in message or "401" in message or "access" in message:
                raise RuntimeError(
                    "Diarization authentication failed. "
                    "Check --hf-token and model access in Hugging Face."
                ) from exc
            raise
        result = assign_word_speakers(
            diarize_segments,
            result,
            fill_nearest=True,
        )

        for idx, seg in enumerate(result.get("segments", []), start=1):
            segments.append(
                {
                    "id": idx,
                    "start": float(seg.get("start", 0.0)),
                    "end": float(seg.get("end", 0.0)),
                    "speaker": seg.get("speaker", "SPEAKER_UNKNOWN"),
                    "text": str(seg.get("text", "")).strip(),
                }
            )
    else:
        # Faster-whisper route: maximum speed for plain transcription.
        model = WhisperModel(
            model_name,
            device="cuda",
            compute_type=compute_type,
            device_index=0,
        )

        backend_device = getattr(model.model, "device", "unknown")
        if str(backend_device) != "cuda":
            raise RuntimeError(
                f"Model backend is not in CUDA (detected: {backend_device}). "
                "CPU fallback is not allowed."
            )

        clip_timestamps = build_clip_timestamps(duration_s, chosen_chunk)
        print(f"Plain ASR -> clip_count={len(clip_timestamps)}")
        batched_model = BatchedInferencePipeline(model=model)
        candidates = batch_candidates(batch_size) if auto_batch else [max(1, batch_size)]

        for candidate_bs in candidates:
            try:
                print(f"Trying batch_size={candidate_bs}")
                segments_iter, _info = batched_model.transcribe(
                    str(audio_file),
                    language=language,
                    log_progress=log_progress,
                    beam_size=beam_size,
                    best_of=best_of,
                    patience=patience,
                    temperature=temperature,
                    condition_on_previous_text=condition_on_previous_text,
                    chunk_length=chosen_chunk,
                    vad_filter=False,
                    clip_timestamps=clip_timestamps,
                    batch_size=candidate_bs,
                    without_timestamps=True,
                )

                segments = []
                for idx, seg in enumerate(segments_iter, start=1):
                    segments.append(
                        {
                            "id": idx,
                            "start": float(seg.start),
                            "end": float(seg.end),
                            "text": seg.text.strip(),
                        }
                    )
                print(f"batch_size={candidate_bs} OK")
                used_batch_size = candidate_bs
                break
            except Exception as exc:
                last_exc = exc
                message = str(exc).lower()
                if "out of memory" in message or ("cuda" in message and "memory" in message):
                    print(f"OOM with batch_size={candidate_bs}, trying lower...")
                    continue
                raise
        else:
            if last_exc is not None:
                raise RuntimeError("All tested batch sizes failed.") from last_exc
            raise RuntimeError("Could not run transcription.")

    elapsed = time.perf_counter() - started
    x_realtime = (duration_s / elapsed) if elapsed > 0 else 0.0

    stem = audio_file.stem
    out_file = output_dir / f"{stem}.{output_format}"

    if output_format == "txt":
        write_txt(segments, out_file)
    elif output_format == "docx":
        title = docx_title or f"Transcripcion - {stem}"
        write_docx(segments, out_file, title=title)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")

    print("Generated file:")
    print(f"- {out_file}")
    print(
        "Performance -> "
        f"elapsed={elapsed:.1f}s, "
        f"x_realtime={x_realtime:.2f}x, "
        f"segments={len(segments)}, "
        f"batch_size={used_batch_size}"
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Local transcription with optional WhisperX diarization and TXT/DOCX output."
    )
    parser.add_argument(
        "--audio-file",
        required=True,
        help="Path to input audio file (wav/mp3/m4a/etc).",
    )
    parser.add_argument(
        "--output-dir",
        default="outputs",
        help="Directory for output files (default: outputs).",
    )
    parser.add_argument(
        "--model",
        default="large-v3",
        help="Model name (default: large-v3).",
    )
    parser.add_argument(
        "--language",
        default="es",
        help="Language code (default: es). Use 'auto' for auto-detection.",
    )
    parser.add_argument(
        "--preset",
        choices=["fast", "balanced", "quality"],
        default="balanced",
        help="Preset profile (default: balanced).",
    )
    parser.add_argument(
        "--batch-size",
        type=int,
        default=32,
        help="Target batch size for batched GPU inference (default: 32).",
    )
    parser.add_argument(
        "--beam-size",
        type=int,
        default=3,
        help="Beam size (default: 3).",
    )
    parser.add_argument(
        "--best-of",
        type=int,
        default=3,
        help="Best-of candidates (default: 3).",
    )
    parser.add_argument(
        "--patience",
        type=float,
        default=1.0,
        help="Decoder patience (default: 1.0).",
    )
    parser.add_argument(
        "--temperature",
        type=float,
        default=0.0,
        help="Decoding temperature (default: 0.0).",
    )
    parser.add_argument(
        "--condition-on-previous-text",
        action="store_true",
        help="Use previous text as context between windows.",
    )
    parser.add_argument(
        "--chunk-length",
        type=int,
        default=30,
        help="Chunk length in seconds (default: 30).",
    )
    parser.add_argument(
        "--log-progress",
        action="store_true",
        help="Show model progress logs.",
    )
    parser.add_argument(
        "--no-auto-batch",
        action="store_true",
        help="Disable automatic batch-size fallback on OOM.",
    )
    parser.add_argument(
        "--compute-type",
        default="float16",
        help="Requested compute type (default: float16).",
    )
    parser.add_argument(
        "--output-format",
        choices=["txt", "docx"],
        default="txt",
        help="Output format (default: txt).",
    )
    parser.add_argument(
        "--docx-title",
        default=None,
        help="Optional title for DOCX output.",
    )
    parser.add_argument(
        "--diarize",
        action="store_true",
        help="Enable speaker diarization (who spoke when) using WhisperX + pyannote.",
    )
    parser.add_argument(
        "--hf-token",
        default=None,
        help="Hugging Face token for diarization models (or set HF_TOKEN env var).",
    )
    parser.add_argument(
        "--diarization-device",
        choices=["cpu", "cuda"],
        default="cpu",
        help="Device for diarization stage (default: cpu).",
    )
    parser.add_argument(
        "--num-speakers",
        type=int,
        default=None,
        help="Known exact number of speakers.",
    )
    parser.add_argument(
        "--min-speakers",
        type=int,
        default=None,
        help="Lower bound for number of speakers.",
    )
    parser.add_argument(
        "--max-speakers",
        type=int,
        default=None,
        help="Upper bound for number of speakers.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    apply_preset(args)
    language = None if args.language.lower() == "auto" else args.language

    try:
        run(
            audio_file=Path(args.audio_file),
            output_dir=Path(args.output_dir),
            model_name=args.model,
            language=language,
            batch_size=args.batch_size,
            requested_compute_type=args.compute_type,
            beam_size=args.beam_size,
            best_of=args.best_of,
            patience=args.patience,
            temperature=args.temperature,
            condition_on_previous_text=args.condition_on_previous_text,
            chunk_length=args.chunk_length,
            log_progress=args.log_progress,
            auto_batch=not args.no_auto_batch,
            output_format=args.output_format,
            docx_title=args.docx_title,
            diarize=args.diarize,
            hf_token=args.hf_token,
            diarization_device=args.diarization_device,
            num_speakers=args.num_speakers,
            min_speakers=args.min_speakers,
            max_speakers=args.max_speakers,
        )
    except RuntimeError as exc:
        raise SystemExit(f"ERROR: {exc}") from exc


if __name__ == "__main__":
    main()
